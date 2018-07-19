import codecs
import collections
import csv
import json
import math
import os
import sys

from pathlib import Path
from pprint import pprint

import jinja2

from docx import Document

from Plant import Plant

INPUT_DIR = Path('1.5')
OUTPUT_DIR = Path('generated/1.5')


def deduplicate_csv():
    csv_file = INPUT_DIR / 'plant_locations_rarities.csv'
    new_csv = OUTPUT_DIR / 'plant_locations_rarities_v2.csv'

    tmp_list = []

    with open(csv_file, 'r') as f:
        reader = csv.reader(f)
        with open(new_csv, 'w', newline='') as new:
            fieldnames = ['Plant',
                          'Arctic',
                          'City/Urban',
                          'Coastal',
                          'Desert',
                          'Forest',
                          'Jungle',
                          'Mountain',
                          'Ocean',
                          'Plain',
                          'River',
                          'Swamp',
                          'Underdark/Cave',
                          'Other',
                          'Rarity']
            writer = csv.DictWriter(new, fieldnames)

            for row in reader:
                name = row[0]
                if name == 'Plant':
                    continue
                rarity = row[14]

                row_regions = {
                    'Arctic': row[1],
                    'City': row[2],
                    'Coastal': row[3],
                    'Desert': row[4],
                    'Forest': row[5],
                    'Jungle': row[6],
                    'Mountain': row[7],
                    'Ocean': row[8],
                    'Plain': row[9],
                    'River': row[10],
                    'Swamp': row[11],
                    'Underdark': row[12],
                    'Other': row[13],
                }

                regions = []
                for region, found in row_regions.items():
                    if found.upper() == 'X':
                        regions.append(region)

                plant = Plant(name, regions, rarity=rarity)

                for tmp_plant in tmp_list:
                    if plant.name == tmp_plant.name:
                        plant.regions.extend(tmp_plant.regions)
                        tmp_list.remove(tmp_plant)

                tmp_list.append(plant)
                # print(plant)
            pprint(tmp_list)
            print('plants:', len(tmp_list))

            writer.writeheader()
            for plant in tmp_list:
                foo = {'Plant': plant.name,
                       'Arctic': 'X' if 'Arctic' in plant.regions else '',
                       'City/Urban': 'X' if 'City' in plant.regions else '',
                       'Coastal': 'X' if 'Coastal' in plant.regions else '',
                       'Desert': 'X' if 'Desert' in plant.regions else '',
                       'Forest': 'X' if 'Forest' in plant.regions else '',
                       'Jungle': 'X' if 'Jungle' in plant.regions else '',
                       'Mountain': 'X' if 'Mountain' in plant.regions else '',
                       'Ocean': 'X' if 'Ocean' in plant.regions else '',
                       'Plain': 'X' if 'Plain' in plant.regions else '',
                       'River': 'X' if 'River' in plant.regions else '',
                       'Swamp': 'X' if 'Swamp' in plant.regions else '',
                       'Underdark/Cave': 'X' if 'Underdark' in plant.regions else '',
                       'Other': 'X' if 'Other' in plant.regions else '',
                       'Rarity': plant.rarity}
                writer.writerow(foo)


def convert_csv_to_json():
    csv_file = OUTPUT_DIR / 'plant_locations_rarities_v2.csv'
    json_file = OUTPUT_DIR / 'plant_info.json'

    with open(csv_file, 'r') as f:
        reader = csv.DictReader(f)
        with open(json_file, 'w') as outfile:

            data = {}

            for row in reader:
                print(row['Plant'])

                row_regions = {
                    'Arctic': row['Arctic'],
                    'City': row['City/Urban'],
                    'Coastal': row['Coastal'],
                    'Desert': row['Desert'],
                    'Forest': row['Forest'],
                    'Jungle': row['Jungle'],
                    'Mountain': row['Mountain'],
                    'Ocean': row['Ocean'],
                    'Plain': row['Plain'],
                    'River': row['River'],
                    'Swamp': row['Swamp'],
                    'Underdark': row['Underdark/Cave'],
                    'Other': row['Other'],
                }

                regions = []
                for region, found in row_regions.items():
                    if found.upper() == 'X':
                        regions.append(region)

                data[row['Plant']] = {'Regions': regions,
                                      'Rarity': row['Rarity']}

            json.dump(data, outfile, indent=4)


def parse_description():
    plant_info_json = OUTPUT_DIR / 'plant_info.json'
    plant_info_v2_json = OUTPUT_DIR / 'plant_info_v2.json'

    docx_file = Path('1.5/plants_v1.5_orig.docx')
    doc = Document(docx_file)

    # for para in range(0, len(doc.paragraphs)):
    #     print(para, doc.paragraphs[para].text)
    # sys.exit()

    plants_by_letter = {
        'A': [79, 238],  # broken
        'B': [238, 365],
        'C': [365, 478],
        'D': [478, 603],
        'E': [603, 650],
        'F': [650, 733],
        'G': [733, 825],
        'H': [825, 923],  # broken
        'I': [923, 949],
        'J': [949, 978],
        'K': [978, 1028],
        'L': [1028, 1103],  # broken
        'M': [1103, 1210],  # broken
        'N': [1210, 1230],
        'O': [1230, 1271],
        'P': [1271, 1330],
        'Q': [1330, 1347],
        'R': [1347, 1418],
        'S': [1418, 1612],
        'T': [1612, 1713],
        'U': [1713, 1724],
        'V': [1724, 1756],
        'W': [1756, 1836],
        'X': [1836, 1841],
        'Y': [1841, 1867],
        'Z': [1867, 1880]
    }

    # Key:   found in word document
    # Value: found in csv (from Excel sheet)
    inconsistencies = {
        "Athelas (Kingsfoil)": "Athelas",
        "Dainaberry (Sleepberry)": "Dainaberry",
        "Devilweed (Wyssin)": "Devilweed",
        "Dungeon Fungus (Dungus)": "Dungeon Fungus",
        "Floure-du-luce (Blue Flag)": "Floure-du-luce",
        "Harrada Leaf (Barbarian's Boon)": "Harrada Leaf",
        "Hathil": "Hathlil",
        "Ipt (Roanwood)": "Ipt",
        "Jelly Moss (Ogre Snot)": "Jelly Moss",
        "Kiss of Discord (Lusiri Blossom)": "Kiss of Discord",
        "Kylathar (Changeberry)": "Kylathar",
        "Laishaberries (Fruit of Silence)": "Laishaberries",
        "Luurden (Bloodfruit)": "Luurden",
        "Maiden's Hair (Earthsilk)": "Maiden's Hair",
        "Obaddis Leaf\t": "Obaddis Leaf",
        "Powdered Desert Milk (Wolves' Milk)": "Powdered Desert Milk",
        "Saracen's Confound": "Saracen's Compound",
        "Silverthorn (Aelebera)": "Silverthorn",
        "Sussur (Deeproot)": "Sussur",
        "Tahtoalethi (Wishfern)": "Tahtoalethi",
        "Tamariske (Blacktear)": "Tamariske",
        "Wildwood (Saelas)": "Wildwood",
        "Yarpick (Daggerthorn)": "Yarpick"
    }

    invalid_names = ['', '\n', '\n ']

    with open(plant_info_json) as json_file:
        data = json.load(json_file)

        # letters that have all normal entries
        for letter, entries in plants_by_letter.items():

            # letters with non-standard size entries
            if letter == 'A':
                print('*****', letter, '*****')
                for para in range(entries[0] + 1, entries[1], 3):
                    if para > 110:
                        para = para + 11
                    if para > 237:
                        continue
                    name = doc.paragraphs[para].text

                    name = name.replace('’', "'")
                    if name in invalid_names:
                        continue
                    if name in inconsistencies.keys():
                        name = inconsistencies[name]

                    # plant = Plant(name, regions, description=description)
                    # regions = doc.paragraphs[para + 1].text  # skipping since getting from csv
                    description = doc.paragraphs[para + 2].text

                    # print(plant, plant.description)
                    data[name]['Description'] = description
                    if name == 'Alil':
                        extra_info = []
                        for line in range(3, 14):
                            extra_info.append(doc.paragraphs[para + line].text)
                        data[name]['Extra info'] = extra_info

                    print(name, data[name])
                continue

            if letter == 'H':
                print('*****', letter, '*****')
                for para in range(entries[0] + 1, entries[1], 3):
                    if para > 838:
                        para = para + 12
                    if para > 921:
                        continue
                    name = doc.paragraphs[para].text

                    name = name.replace('’', "'")
                    if name in invalid_names:
                        continue
                    if name in inconsistencies.keys():
                        name = inconsistencies[name]

                    # plant = Plant(name, regions, description=description)
                    # regions = doc.paragraphs[para + 1].text  # skipping since getting from csv
                    description = doc.paragraphs[para + 2].text

                    # print(plant, plant.description)
                    data[name]['Description'] = description
                    if name == 'Hangman Tree':
                        extra_info = []
                        for line in range(3, 15):
                            extra_info.append(doc.paragraphs[para + line].text)
                        data[name]['Extra info'] = extra_info

                    print(name, data[name])
                continue

            if letter == 'L':
                print('*****', letter, '*****')
                for para in range(entries[0] + 1, entries[1], 3):
                    if para > 1077:
                        para = para + 4
                    if para > 1101:
                        continue
                    name = doc.paragraphs[para].text

                    name = name.replace('’', "'")
                    if name in invalid_names:
                        continue
                    if name in inconsistencies.keys():
                        name = inconsistencies[name]

                    # plant = Plant(name, regions, description=description)
                    # regions = doc.paragraphs[para + 1].text  # skipping since getting from csv
                    description = doc.paragraphs[para + 2].text

                    # print(plant, plant.description)
                    data[name]['Description'] = description
                    if name == 'Lizuara':
                        extra_info = []
                        for line in range(3, 7):
                            extra_info.append(doc.paragraphs[para + line].text)
                        data[name]['Extra info'] = extra_info

                    print(name, data[name])
                continue

            if letter == 'M':
                print('*****', letter, '*****')
                for para in range(entries[0] + 1, entries[1], 3):
                    if para > 1134:
                        para = para + 12
                    if para > 1208:
                        continue
                    name = doc.paragraphs[para].text

                    name = name.replace('’', "'")
                    if name in invalid_names:
                        continue
                    if name in inconsistencies.keys():
                        name = inconsistencies[name]

                    # plant = Plant(name, regions, description=description)
                    # regions = doc.paragraphs[para + 1].text  # skipping since getting from csv
                    description = doc.paragraphs[para + 2].text

                    # print(plant, plant.description)
                    data[name]['Description'] = description
                    if name == 'Marsh Maw':
                        extra_info = []
                        for line in range(3, 15):
                            extra_info.append(doc.paragraphs[para + line].text)
                        data[name]['Extra info'] = extra_info

                    print(name, data[name])
                continue

            # letters with standard size entries
            print('*****', letter, '*****')
            for para in range(entries[0] + 1, entries[1], 3):
                name = doc.paragraphs[para].text

                if name in invalid_names:
                    continue
                if '’' in name:
                    name = name.replace('’', "'")
                if name in inconsistencies.keys():
                    name = inconsistencies[name]

                regions = doc.paragraphs[para + 1].text
                description = doc.paragraphs[para + 2].text

                # plant = Plant(name, regions, description=description)
                # print(plant, plant.description)
                data[name]['Description'] = description
                print(name, data[name])

                # if len(doc.paragraphs[para].text) == 1 and doc.paragraphs[para].text.isalpha():
                #     print(para, doc.paragraphs[para].text)

        with open(plant_info_v2_json, 'w') as outfile:
            json.dump(data, outfile, indent=4)


def parse_rarity():
    plant_info_v2_json = OUTPUT_DIR / 'plant_info_v2.json'
    plant_info_v3_json = OUTPUT_DIR / 'plant_info_v3.json'

    docx_file = Path('1.5/plants_v1.5_orig.docx')
    doc = Document(docx_file)

    plants_by_rarity = {
        'Very Common': [2958, 3117],
        'Common': [3118, 3267],
        'Uncommon': [3268, 3411],
        'Rare': [3412, 3535],
        'Very Rare': [3536, 3630],
        'Legendary': [3632, 3681]
    }

    # Key:   found in word document
    # Value: found in csv (from Excel sheet)
    inconsistencies = {
        "Airplant": "Air Plant",
        "Bloostaunch": "Bloodstaunch",
        "Calcena Mushrooms": "Calcena Mushroom",
        "Cow Parnsip": "Cow Parsnip",
        "Saracen's Confound": "Saracen's Compound",
        "Silver Hisbiscus": "Silver Hibiscus"
    }
    invalid_names = ['', '\n', '\n ',
                     'L ', 'P\t', 'P (Cont.)', 'S (Cont.)',
                     'Very Common', 'Common', 'Uncommon', 'Rare', 'Very Rare ', 'Legendary']

    with open(plant_info_v2_json) as json_file:
        data = json.load(json_file)
        for rarity, entries in plants_by_rarity.items():
            print('*****', rarity, '*****')
            for para in range(entries[0], entries[1]):
                name = doc.paragraphs[para].text

                name = name.replace('’', "'")
                if name in invalid_names:
                    continue
                if len(name) == 1:
                    continue
                if name in inconsistencies.keys():
                    name = inconsistencies[name]

                data[name]['Rarity'] = rarity
                print(name, data[name])

        with open(plant_info_v3_json, 'w') as outfile:
            json.dump(data, outfile, indent=4)


def generate_homebrewery_markdown():
    plant_info_json = OUTPUT_DIR / 'plant_info_v3.json'
    homebrew_file = OUTPUT_DIR / "generated-homebrew.txt"

    rarity_symbols = {
        'Very Common': 'VC',
        'Common': 'C',
        'Uncommon': 'U',
        'Rare': 'R',
        'Very Rare': 'VR',
        'Legendary': 'L'
    }

    plants_by_letter = {}
    plants_by_region = collections.OrderedDict({
        'Arctic': [],     # D&D 5E standard
        'City': [],
        'Coastal': [],    # D&D 5E standard - Coast
        'Desert': [],     # D&D 5E standard
        'Forest': [],     # D&D 5E standard
        'Jungle': [],
        'Plain': [],      # D&D 5E standard - Grassland
        'Mountain': [],   # D&D 5E standard
        'Ocean': [],
        'River': [],
        'Swamp': [],      # D&D 5E standard
        'Underdark': [],  # D&D 5E standard
        'Other': []
    })
    plants_by_rarity = collections.OrderedDict({
        'Very Common': [],
        'Common': [],
        'Uncommon': [],
        'Rare': [],
        'Very Rare': [],
        'Legendary': []
    })

    pages_before_plant_entries = 6
    header_height = 5  # equivalent # of lines for plant name, location, rarity
    footer_height = 1  # equivalent # of lines after entry
    lines_available_per_column = 60
    page_height = lines_available_per_column * 2
    desc_line_length = 60

    with open(plant_info_json) as json_file:
        data = json.load(json_file)

        pages = {}
        page_num = 1
        pages[page_num] = []
        current_page_height = 0

        for plant, entry in data.items():
            # plant = name
            # entry = {
            #   Regions: [x, y, z]
            #   Rarity: "xyz"
            #   Description: "foo"
            #   Extra info: ["this", "is", "extra"]
            #   ### added below ###
            #   name: "blah"
            #   page: 34
            #   first_letter: "B"
            #   homebrewery_height: 12
            # }

            first_letter = plant[0].upper()
            description_height = math.ceil(len(entry['Description']) / desc_line_length)
            homebrewery_height = header_height + description_height + footer_height
            if plant == 'Alil':
                homebrewery_height += 40
            if plant == 'Hangman Tree':
                homebrewery_height += 40
            if plant == 'Lizuara':
                homebrewery_height += 15
            if plant == 'Marsh Maw':
                homebrewery_height += 50
            current_page_height += homebrewery_height

            if current_page_height > page_height:
                # print('page:', page_num, '   height:', current_page_height)
                current_page_height = homebrewery_height
                page_num += 1
                pages[page_num] = []

            entry['name'] = plant
            entry['page'] = page_num + pages_before_plant_entries
            entry['first_letter'] = first_letter
            entry['homebrewery_height'] = homebrewery_height
            entry['rarity_symbol'] = rarity_symbols[entry['Rarity']]

            # Manually give descriptions line breaks to make content fit on page
            if plant == 'Darmanzar Stalk':
                entry['Description'] = entry['Description'].replace('Living characters', '\n\nLiving characters')
            if plant == 'Hidden Hibiscus':
                entry['Description'] = entry['Description'].replace('Bloom colors', '\n\nBloom colors')

            # manually deal with the 4 "special" plants
            if plant == 'Alil':
                entry['Extra info'] = '''| d10 | Psionic Ability |
|:----:|:-------------|
| 1  | Temporary intelligence bonus of 1. |
| 2  | Precision Mind: you become skilled at reading your foes. With this capability, you can call upon your gift to strike with increased accuracy. You add your Intelligence modifier (minimum 1) to an attack roll, after the roll, but before the announcement of the result. You can use this feature a number of times equal to your Intelligence modifier (minimum of 1). |
| 3  | Immunity to psychic damage. |
| 4  | You cannot be charmed or frightened for the duration. |
| 5  | You may take two actions during each of your turns for the duration. |
| 6  | You may add 1d6 psychic damage to any attack you make of 5 or less damage. |
| 7  | You project a field of improbability around yourself, creating a fleeting protective shell, gaining  +4 temporary bonus to AC. |
| 8  | You may, as a bonus action once per combat, instantly deliver a massive assault on the thought pathways of any one creature, dealing 1d10 points of damage to it. |
| 9  | You heal another creature\u2019s wounds, transferring some of its damage to yourself. When you manifest this power, you can heal as much as 2d10 points of damage to an ally by taking half of this damage to yourself. |
| 10 | You may gain proficiency in one extra skill until the long rest or proficiency with one tool or instrument permanently. |
            '''

            if plant == 'Hangman Tree':
                entry['Description'] = entry['Description'].replace('The main body', '\n\nThe main body')
                entry['Extra info'] = '''___
> ## Hangman Tree
>*Large tree, unaligned*
> ___
> - **Armor Class** 12
> - **Hit Points** 172(15d12 + 80)
> - **Speed** 1ft.
>___
>|STR|DEX|CON|INT|WIS|CHA|
>|:---:|:---:|:---:|:---:|:---:|:---:|
>|20 (+5)|6 (-2)|20 (+5)|3 (-4)|6 (-2)|4 (-3)|
>___
> - **Damage Resistances** bludgeoning, piercing
> - **Damage Vulnerabilities** fire
> - **Condition Immunities** blinded, deafened, frightened, exhaustion
> - **Senses** blindsight 90 ft., passive Perception 8
> - **Challenge** 8 (3,800 XP)
> ___

> ### Actions
> ***Multiattack.*** The hangman tree may make three *constrict* attacks per turn.
>
> ***Constrict.*** *Melee Weapon Attack:* +9 to hit, reach 10 ft., one target. *Hit* 11 (1d8 + 5) bludgeoning damage, and a Large or smaller target is grappled (escape DC 16).
            '''

            if plant == 'Lizuara':
                extra_info = ''
                for line in entry['Extra info']:
                    extra_info = extra_info + '\n\n' + line
                entry['Extra info'] = extra_info

            if plant == 'Marsh Maw':
                entry['Extra info'] = '''___
> ## Marsh Maw
>*Large plant, unaligned*
> ___
> - **Armor Class** 14
> - **Hit Points** 112 (14d10 + 28)
> - **Speed** 20 ft.. swim 20 ft.
>___
>|STR|DEX|CON|INT|WIS|CHA|
>|:---:|:---:|:---:|:---:|:---:|:---:|
>|16 (+3)|12 (+1)|14 (+2)|4 (-3)|10 (+0)|6 (-2)|
>___
> - **Damage Resistances** bludgeoning
> - **Condition Immunities** blinded, deafened, frightened, prone
> - **Senses** blindsight 60 ft., passive Perception 10
> - **Challenge** 4 (1,100 XP)
> ___
> ### Actions
> ***Multiattack.*** The marsh maw can make two *constrict* attacks and a *bite* attack.
>
> ***Constrict.*** *Melee Weapon Attack:* +5 to hit, reach 10 ft., one target. *Hit* 11 (2d6 + 3) bludgeoning damage, and a Large or smaller target is grappled (escape DC 14). Until this grapple ends, the target is restrained, and the marsh maw cannot constrict another target.
>
> ***Bite.*** *Melee Weapon Attack:* +5 to hit, reach 5 ft., one Medium or smaller target. Hit: 11 (2d6+3) piercing damage, and a target is blinded, restrained, and unable to breathe. The target must succeed on a DC 14 Cons. Save at the start of each of the marsh maw’s turns or take 11 (2d8+3) bludgeoning damage. If the marsh maw moves, the engulfed target moves with it. The marsh maw is unable to use the bite attack until it releases the held creature.
                '''

            pages[page_num].append(entry)

            if first_letter in plants_by_letter:
                plants_by_letter[first_letter].append(entry)
            else:
                plants_by_letter[first_letter] = [entry]

            rarity = entry['Rarity']
            if rarity in plants_by_rarity:
                plants_by_rarity[rarity].append(entry)

            regions = entry['Regions']
            for region in regions:
                if region in plants_by_region:
                    plants_by_region[region].append(entry)
                else:
                    plants_by_region[region] = [entry]

        plants_by_rarity = collections.OrderedDict(sorted(plants_by_rarity.items(), key=lambda t: len(t)))

        context = {
            'title': 'Broderick’s Compendium of Fantasy Plants',
            'plants_by_letter': plants_by_letter,
            'plants_by_region': plants_by_region,
            'plants_by_rarity': plants_by_rarity,
            'pages': pages
        }
        result = render('templates/homebrew.md.j2', context)

        # print(result)
        with codecs.open(homebrew_file, "w", encoding="utf-8") as hb_page:
            hb_page.write(result)
        # with open('generated-homebrew.txt', 'w') as hb_page:
        #     hb_page.write(result)


def render(tpl_path, context):
    path, filename = os.path.split(tpl_path)
    env = jinja2.Environment(loader=jinja2.FileSystemLoader(path or './'))
    env.trim_blocks = True
    env.lstrip_blocks = True
    return env.get_template(filename).render(context)


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    # deduplicate_csv()
    # convert_csv_to_json()
    parse_description()
    # parse_rarity()
    # generate_homebrewery_markdown()


if __name__ == "__main__":
    main()
