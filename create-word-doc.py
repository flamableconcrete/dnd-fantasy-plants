import csv
import json
import math
import os
import sys

from pprint import pprint

import jinja2

from docx import Document


class Plant(object):

    def __init__(self, name, regions, rarity='', description=''):
        self.name = name
        self.regions = regions
        self.rarity = rarity
        self.description = description

        foo = '{}.json'.format(self.name).lower().replace(' ', '_')
        self.filename = ''.join(ch for ch in foo if ch.isalnum() or ch == '.' or ch == '_')

    def __repr__(self):
        return '{} ({})'.format(self.name, self.regions)


def deduplicate_csv():
    csv_file = 'plant_locations_rarities.csv'
    new_csv = 'plant_locations_rarities_v2.csv'

    tmp_list = []

    with open(csv_file, 'r') as f:
        reader = csv.reader(f)
        with open(new_csv, 'w', newline='') as new:
            fieldnames = ['Plant',
                          'Arctic',
                          'City/Urban',
                          'Coastal',
                          'Desert',
                          'Forest',
                          'Jungle',
                          'Mountain',
                          'Ocean',
                          'Plain',
                          'River',
                          'Swamp',
                          'Underdark/Cave',
                          'Other',
                          'Rarity']
            writer = csv.DictWriter(new, fieldnames)

            for row in reader:
                name = row[0]
                if name == 'Plant':
                    continue
                rarity = row[14]

                row_regions = {
                    'Arctic': row[1],
                    'City': row[2],
                    'Coastal': row[3],
                    'Desert': row[4],
                    'Forest': row[5],
                    'Jungle': row[6],
                    'Mountain': row[7],
                    'Ocean': row[8],
                    'Plain': row[9],
                    'River': row[10],
                    'Swamp': row[11],
                    'Underdark': row[12],
                    'Other': row[13],
                }

                regions = []
                for region, found in row_regions.items():
                    if found.upper() == 'X':
                        regions.append(region)

                plant = Plant(name, regions, rarity=rarity)

                for tmp_plant in tmp_list:
                    if plant.name == tmp_plant.name:
                        plant.regions.extend(tmp_plant.regions)
                        tmp_list.remove(tmp_plant)

                tmp_list.append(plant)
                # print(plant)
            pprint(tmp_list)
            print('plants:', len(tmp_list))

            writer.writeheader()
            for plant in tmp_list:
                foo = {'Plant': plant.name,
                       'Arctic': 'X' if 'Arctic' in plant.regions else '',
                       'City/Urban': 'X' if 'City' in plant.regions else '',
                       'Coastal': 'X' if 'Coastal' in plant.regions else '',
                       'Desert': 'X' if 'Desert' in plant.regions else '',
                       'Forest': 'X' if 'Forest' in plant.regions else '',
                       'Jungle': 'X' if 'Jungle' in plant.regions else '',
                       'Mountain': 'X' if 'Mountain' in plant.regions else '',
                       'Ocean': 'X' if 'Ocean' in plant.regions else '',
                       'Plain': 'X' if 'Plain' in plant.regions else '',
                       'River': 'X' if 'River' in plant.regions else '',
                       'Swamp': 'X' if 'Swamp' in plant.regions else '',
                       'Underdark/Cave': 'X' if 'Underdark' in plant.regions else '',
                       'Other': 'X' if 'Other' in plant.regions else '',
                       'Rarity': plant.rarity}
                writer.writerow(foo)


def convert_csv_to_json():
    csv_file = 'plant_locations_rarities_v2.csv'
    json_file = 'plant_info.json'

    with open(csv_file, 'r') as f:
        reader = csv.DictReader(f)
        with open(json_file, 'w') as outfile:

            data = {}

            for row in reader:
                print(row['Plant'])

                row_regions = {
                    'Arctic': row['Arctic'],
                    'City': row['City/Urban'],
                    'Coastal': row['Coastal'],
                    'Desert': row['Desert'],
                    'Forest': row['Forest'],
                    'Jungle': row['Jungle'],
                    'Mountain': row['Mountain'],
                    'Ocean': row['Ocean'],
                    'Plain': row['Plain'],
                    'River': row['River'],
                    'Swamp': row['Swamp'],
                    'Underdark': row['Underdark/Cave'],
                    'Other': row['Other'],
                }

                regions = []
                for region, found in row_regions.items():
                    if found.upper() == 'X':
                        regions.append(region)

                data[row['Plant']] = {'Regions': regions,
                                      'Rarity': row['Rarity']}

            json.dump(data, outfile, indent=4)


def parse_description():
    plant_info_json = 'plant_info.json'
    plant_info_v2_json = 'plant_info_v2.json'

    docx_file = 'plants_v1.5_orig.docx'
    doc = Document(docx_file)

    # for para in range(0, len(doc.paragraphs)):
    #     print(para, doc.paragraphs[para].text)
    # sys.exit()

    plants_by_letter = {
        'A': [79, 238],  # broken
        'B': [238, 365],
        'C': [365, 478],
        'D': [478, 603],
        'E': [603, 650],
        'F': [650, 733],
        'G': [733, 825],
        'H': [825, 923],  # broken
        'I': [923, 949],
        'J': [949, 978],
        'K': [978, 1028],
        'L': [1028, 1103],  # broken
        'M': [1103, 1210],  # broken
        'N': [1210, 1230],
        'O': [1230, 1271],
        'P': [1271, 1330],
        'Q': [1330, 1347],
        'R': [1347, 1418],
        'S': [1418, 1612],
        'T': [1612, 1713],
        'U': [1713, 1724],
        'V': [1724, 1756],
        'W': [1756, 1836],
        'X': [1836, 1841],
        'Y': [1841, 1867],
        'Z': [1867, 1880]
    }

    # Key:   found in word document
    # Value: found in csv (from Excel sheet)
    inconsistencies = {
        "Athelas (Kingsfoil)": "Athelas",
        "Dainaberry (Sleepberry)": "Dainaberry",
        "Devilweed (Wyssin)": "Devilweed",
        "Dungeon Fungus (Dungus)": "Dungeon Fungus",
        "Floure-du-luce (Blue Flag)": "Floure-du-luce",
        "Harrada Leaf (Barbarian's Boon)": "Harrada Leaf",
        "Hathil": "Hathlil",
        "Ipt (Roanwood)": "Ipt",
        "Jelly Moss (Ogre Snot)": "Jelly Moss",
        "Kiss of Discord (Lusiri Blossom)": "Kiss of Discord",
        "Kylathar (Changeberry)": "Kylathar",
        "Laishaberries (Fruit of Silence)": "Laishaberries",
        "Luurden (Bloodfruit)": "Luurden",
        "Maiden's Hair (Earthsilk)": "Maiden's Hair",
        "Obaddis Leaf\t": "Obaddis Leaf",
        "Powdered Desert Milk (Wolves' Milk)": "Powdered Desert Milk",
        "Saracen's Confound": "Saracen's Compound",
        "Silverthorn (Aelebera)": "Silverthorn",
        "Sussur (Deeproot)": "Sussur",
        "Tahtoalethi (Wishfern)": "Tahtoalethi",
        "Tamariske (Blacktear)": "Tamariske",
        "Wildwood (Saelas)": "Wildwood",
        "Yarpick (Daggerthorn)": "Yarpick"
    }

    invalid_names = ['', '\n', '\n ']

    with open(plant_info_json) as json_file:
        data = json.load(json_file)

        # letters that have all normal entries
        for letter, entries in plants_by_letter.items():

            # letters with non-standard size entries
            if letter == 'A':
                print('*****', letter, '*****')
                for para in range(entries[0] + 1, entries[1], 3):
                    if para > 110:
                        para = para + 11
                    if para > 237:
                        continue
                    name = doc.paragraphs[para].text

                    name = name.replace('’', "'")
                    if name in invalid_names:
                        continue
                    if name in inconsistencies.keys():
                        name = inconsistencies[name]

                    # plant = Plant(name, regions, description=description)
                    # regions = doc.paragraphs[para + 1].text  # skipping since getting from csv
                    description = doc.paragraphs[para + 2].text

                    # print(plant, plant.description)
                    data[name]['Description'] = description
                    if name == 'Alil':
                        extra_info = []
                        for line in range(3, 14):
                            extra_info.append(doc.paragraphs[para + line].text)
                        data[name]['Extra info'] = extra_info

                    print(name, data[name])
                continue

            if letter == 'H':
                print('*****', letter, '*****')
                for para in range(entries[0] + 1, entries[1], 3):
                    if para > 838:
                        para = para + 12
                    if para > 921:
                        continue
                    name = doc.paragraphs[para].text

                    name = name.replace('’', "'")
                    if name in invalid_names:
                        continue
                    if name in inconsistencies.keys():
                        name = inconsistencies[name]

                    # plant = Plant(name, regions, description=description)
                    # regions = doc.paragraphs[para + 1].text  # skipping since getting from csv
                    description = doc.paragraphs[para + 2].text

                    # print(plant, plant.description)
                    data[name]['Description'] = description
                    if name == 'Hangman Tree':
                        extra_info = []
                        for line in range(3, 15):
                            extra_info.append(doc.paragraphs[para + line].text)
                        data[name]['Extra info'] = extra_info

                    print(name, data[name])
                continue

            if letter == 'L':
                print('*****', letter, '*****')
                for para in range(entries[0] + 1, entries[1], 3):
                    if para > 1077:
                        para = para + 4
                    if para > 1101:
                        continue
                    name = doc.paragraphs[para].text

                    name = name.replace('’', "'")
                    if name in invalid_names:
                        continue
                    if name in inconsistencies.keys():
                        name = inconsistencies[name]

                    # plant = Plant(name, regions, description=description)
                    # regions = doc.paragraphs[para + 1].text  # skipping since getting from csv
                    description = doc.paragraphs[para + 2].text

                    # print(plant, plant.description)
                    data[name]['Description'] = description
                    if name == 'Lizuara':
                        extra_info = []
                        for line in range(3, 7):
                            extra_info.append(doc.paragraphs[para + line].text)
                        data[name]['Extra info'] = extra_info

                    print(name, data[name])
                continue

            if letter == 'M':
                print('*****', letter, '*****')
                for para in range(entries[0] + 1, entries[1], 3):
                    if para > 1134:
                        para = para + 12
                    if para > 1208:
                        continue
                    name = doc.paragraphs[para].text

                    name = name.replace('’', "'")
                    if name in invalid_names:
                        continue
                    if name in inconsistencies.keys():
                        name = inconsistencies[name]

                    # plant = Plant(name, regions, description=description)
                    # regions = doc.paragraphs[para + 1].text  # skipping since getting from csv
                    description = doc.paragraphs[para + 2].text

                    # print(plant, plant.description)
                    data[name]['Description'] = description
                    if name == 'Marsh Maw':
                        extra_info = []
                        for line in range(3, 15):
                            extra_info.append(doc.paragraphs[para + line].text)
                        data[name]['Extra info'] = extra_info

                    print(name, data[name])
                continue

            # letters with standard size entries
            print('*****', letter, '*****')
            for para in range(entries[0] + 1, entries[1], 3):
                name = doc.paragraphs[para].text

                if name in invalid_names:
                    continue
                if '’' in name:
                    name = name.replace('’', "'")
                if name in inconsistencies.keys():
                    name = inconsistencies[name]

                regions = doc.paragraphs[para + 1].text
                description = doc.paragraphs[para + 2].text

                # plant = Plant(name, regions, description=description)
                # print(plant, plant.description)
                data[name]['Description'] = description
                print(name, data[name])

                # if len(doc.paragraphs[para].text) == 1 and doc.paragraphs[para].text.isalpha():
                #     print(para, doc.paragraphs[para].text)

        with open(plant_info_v2_json, 'w') as outfile:
            json.dump(data, outfile, indent=4)


def parse_rarity():
    plant_info_v2_json = 'plant_info_v2.json'
    plant_info_v3_json = 'plant_info_v3.json'

    docx_file = 'plants_v1.5_orig.docx'
    doc = Document(docx_file)

    plants_by_rarity = {
        'Very Common': [2958, 3117],
        'Common': [3118, 3267],
        'Uncommon': [3268, 3411],
        'Rare': [3412, 3535],
        'Very Rare': [3536, 3630],
        'Legendary': [3632, 3681]
    }

    # Key:   found in word document
    # Value: found in csv (from Excel sheet)
    inconsistencies = {
        "Airplant": "Air Plant",
        "Bloostaunch": "Bloodstaunch",
        "Calcena Mushrooms": "Calcena Mushroom",
        "Cow Parnsip": "Cow Parsnip",
        "Saracen's Confound": "Saracen's Compound",
        "Silver Hisbiscus": "Silver Hibiscus"
    }
    invalid_names = ['', '\n', '\n ',
                     'L ', 'P\t', 'P (Cont.)', 'S (Cont.)',
                     'Very Common', 'Common', 'Uncommon', 'Rare', 'Very Rare ', 'Legendary']

    with open(plant_info_v2_json) as json_file:
        data = json.load(json_file)
        for rarity, entries in plants_by_rarity.items():
            print('*****', rarity, '*****')
            for para in range(entries[0], entries[1]):
                name = doc.paragraphs[para].text

                name = name.replace('’', "'")
                if name in invalid_names:
                    continue
                if len(name) == 1:
                    continue
                if name in inconsistencies.keys():
                    name = inconsistencies[name]

                data[name]['Rarity'] = rarity
                print(name, data[name])

        with open(plant_info_v3_json, 'w') as outfile:
            json.dump(data, outfile, indent=4)


def generate_homebrewery_markdown():
    plant_info_json = 'plant_info_v3.json'
    plants_by_letter = {}
    with open(plant_info_json) as json_file:
        data = json.load(json_file)
        for plant, entry in data.items():
            first_letter = plant[0].upper()
            if first_letter in plants_by_letter:
                plants_by_letter[first_letter].append(plant)
            else:
                plants_by_letter[first_letter] = [plant]

    header_height = 4  # equivalent # of lines for plant name, location, rarity
    footer_height = 1  # equivalent # of lines after entry
    lines_available_per_column = 60
    page_height = lines_available_per_column * 2
    desc_line_length = 60

    with open(plant_info_json) as json_file:
        data = json.load(json_file)

        pages = {}
        page_num = 1
        pages[page_num] = []
        current_page_height = 0

        for plant, value in data.items():
            # plant = name
            # value = {
            #   Regions: [x, y, z]
            #   Rarity: "xyz"
            #   Description: "foo"
            # }
            value['name'] = plant
            description_height = math.ceil(len(value['Description']) / desc_line_length)
            homebrewery_height = header_height + description_height + footer_height
            # print('    * ', plant, 'height:', homebrewery_height)
            value['homebrewery_height'] = homebrewery_height

            current_page_height += homebrewery_height

            if current_page_height > page_height:
                print('page:', page_num, '   height:', current_page_height)
                current_page_height = homebrewery_height
                page_num += 1
                pages[page_num] = []

            pages[page_num].append(value)

        context = {
            'title': 'Broderick’s Compendium of Fantasy Plants',
            'letters': plants_by_letter,
            'pages': pages
        }
        result = render('templates/homebrew.md.j2', context)

        # print(result)
        with open('generated-homebrew.txt', 'w') as hb_page:
            hb_page.write(result)


def render(tpl_path, context):
    path, filename = os.path.split(tpl_path)
    env = jinja2.Environment(loader=jinja2.FileSystemLoader(path or './'))
    env.trim_blocks = True
    env.lstrip_blocks = True
    return env.get_template(filename).render(context)


def main():
    # deduplicate_csv()
    # convert_csv_to_json()
    # parse_description()
    # parse_rarity()
    generate_homebrewery_markdown()


if __name__ == "__main__":
    main()
