import codecs
import collections
import csv
import json
import math
import os
import sys

from pathlib import Path
from pprint import pprint
from textwrap import dedent

import jinja2

from docx import Document
from icecream import ic

from Plant import Plant

INPUT_DIR = Path('2.0')
OUTPUT_DIR = Path('generated/2.0')
TEMPLATE_DIR = Path('templates')


def deduplicate_csv():
    csv_file = INPUT_DIR / 'plant_locations_rarities_2.0.csv'
    new_csv = OUTPUT_DIR / 'plant_locations_rarities_2.0_v2.csv'

    tmp_list = []

    with open(csv_file, 'r', encoding='utf-8') as f:
        reader = csv.reader(f)
        with open(new_csv, 'w', newline='') as new:
            fieldnames = ['Plant',
                          'Arctic',
                          'City/Urban',
                          'Coastal',
                          'Desert',
                          'Forest',
                          'Jungle',
                          'Mountain',
                          'Ocean',
                          'Plain',
                          'River',
                          'Swamp',
                          'Underdark/Cave',
                          'Other',
                          'Rarity',
                          'Notes']
            writer = csv.DictWriter(new, fieldnames)

            for row in reader:
                name = row[0].replace('\u2019', "'")
                if 'Plant Name' in name:
                    continue

                if row[1].upper() == 'X':
                    rarity = 'Very Common'
                elif row[2].upper() == 'X':
                    rarity = 'Common'
                elif row[3].upper() == 'X':
                    rarity = 'Uncommon'
                elif row[4].upper() == 'X':
                    rarity = 'Rare'
                elif row[5].upper() == 'X':
                    rarity = 'Very Rare'
                elif row[6].upper() == 'X':
                    rarity = 'Legendary'
                else:
                    rarity = 'Other'

                notes = row[7]
                region = [row[8]]

                plant = Plant(name, region, rarity=rarity, notes=notes)

                for tmp_plant in tmp_list:
                    if plant.name == tmp_plant.name:
                        plant.regions.extend(tmp_plant.regions)
                        tmp_list.remove(tmp_plant)

                tmp_list.append(plant)
                # print(plant)

            # pprint(tmp_list)
            print('plants:', len(tmp_list))

            writer.writeheader()
            for plant in tmp_list:
                # print(f'{plant.name}: {plant.regions}')
                foo = {'Plant': plant.name,
                       'Arctic': 'X' if 'Arctic' in plant.regions else '',
                       'City/Urban': 'X' if 'City/Urban' in plant.regions else '',
                       'Coastal': 'X' if 'Coastal' in plant.regions else '',
                       'Desert': 'X' if 'Desert' in plant.regions else '',
                       'Forest': 'X' if 'Forest' in plant.regions else '',
                       'Jungle': 'X' if 'Jungle' in plant.regions else '',
                       'Mountain': 'X' if 'Mountain' in plant.regions else '',
                       'Ocean': 'X' if 'Ocean' in plant.regions else '',
                       'Plain': 'X' if 'Plain' in plant.regions else '',
                       'River': 'X' if 'River' in plant.regions else '',
                       'Swamp': 'X' if 'Swamp' in plant.regions else '',
                       'Underdark/Cave': 'X' if 'Underdark/Cave' in plant.regions else '',
                       'Other': 'X' if 'Other' in plant.regions else '',
                       'Rarity': plant.rarity,
                       'Notes': plant.notes}
                writer.writerow(foo)


def convert_csv_to_json():
    csv_file = OUTPUT_DIR / 'plant_locations_rarities_2.0_v2.csv'
    json_file = OUTPUT_DIR / 'plant_info.json'

    with open(csv_file, 'r') as f:
        reader = csv.DictReader(f)
        with open(json_file, 'w') as outfile:

            data = {}

            for row in reader:
                # print(row['Plant'])

                row_regions = {
                    'Arctic': row['Arctic'],
                    'City': row['City/Urban'],
                    'Coastal': row['Coastal'],
                    'Desert': row['Desert'],
                    'Forest': row['Forest'],
                    'Jungle': row['Jungle'],
                    'Mountain': row['Mountain'],
                    'Ocean': row['Ocean'],
                    'Plain': row['Plain'],
                    'River': row['River'],
                    'Swamp': row['Swamp'],
                    'Underdark': row['Underdark/Cave'],
                    'Other': row['Other'],
                    'Notes': row['Notes'],
                }

                regions = []
                for region, found in row_regions.items():
                    if found.upper() == 'X':
                        regions.append(region)

                data[row['Plant']] = {'Regions': regions,
                                      'Rarity': row['Rarity'],
                                      'Notes': row['Notes']}

            json.dump(data, outfile, indent=4)


def clean_name(name, invalid_names, inconsistencies):
    name = name.replace('’', "'")
    if name in invalid_names:
        return None
    if name in inconsistencies.keys():
        name = inconsistencies[name]
    return name


def process_letter_entries(letter, entries, doc, invalid_names, inconsistencies, data,
                           stupid_entry=None, se_start=None, se_len_extra_info=None, letter_end=None):
    print('*****', letter, '*****')
    for para in range(entries[0] + 1, entries[1], 3):
        if stupid_entry:
            if para > se_start:
                para = para + se_len_extra_info
                # Midnight Coneflower - I hate you for being the only double stupid entry in a letter section
                if letter == 'M' and para > 1342:
                    para = para + 2
            if para > letter_end:
                continue
        name = clean_name(doc.paragraphs[para].text, invalid_names, inconsistencies)
        if not name:
            continue

        # regions = doc.paragraphs[para + 1].text
        description = doc.paragraphs[para + 2].text

        # plant = Plant(name, regions, description=description)
        # print(plant, plant.description)
        data[name]['Description'] = description

        if stupid_entry:
            if name == stupid_entry:
                extra_info = []
                for line in range(3, 3 + se_len_extra_info):
                    extra_info.append(doc.paragraphs[para + line].text)
                data[name]['Extra info'] = extra_info

        print(name, f'p{para}', data[name])

    return data


def parse_description():
    plant_info_json = OUTPUT_DIR / 'plant_info.json'
    plant_info_v2_json = OUTPUT_DIR / 'plant_info_v2.json'

    docx_file = INPUT_DIR / 'Plants_v2.0.docx'
    doc = Document(docx_file)

    # with open('foo.txt', 'w') as f:
    #     for para in range(0, len(doc.paragraphs)):
    #         print(para, doc.paragraphs[para].text, file=f)
    # sys.exit()

    plants_by_letter = {
        'A': [105, 297],  # broken
        'B': [297, 445],
        'C': [445, 581],  # broken
        'D': [581, 718],
        'E': [718, 777],
        'F': [777, 872],
        'G': [872, 967],
        'H': [967, 1068],  # broken
        'I': [1068, 1094],
        'J': [1094, 1123],
        'K': [1123, 1179],
        'L': [1179, 1268],  # broken
        'M': [1268, 1399],  # broken
        'N': [1399, 1434],
        'O': [1434, 1475],
        'P': [1475, 1537],
        'Q': [1537, 1554],
        'R': [1554, 1628],
        'S': [1628, 1838],
        'T': [1838, 1948],
        'U': [1948, 1959],
        'V': [1959, 1994],
        'W': [1994, 2074],
        'X': [2074, 2079],
        'Y': [2079, 2112],
        'Z': [2112, 2125]
    }

    # Key:   found in word document
    # Value: found in csv (from Excel sheet)
    inconsistencies = {
        "Athelas (Kingsfoil)": "Athelas",
        "Alchemilla (Lady's Mantle)": "Alchemilla",
        "All-Heale (Attorlaðe)": "All-Heale",
        "Bodhi Tree (Sacred Fig)": "Bodhi Tree",
        "Dahkra (Dognap)": "Dahkra",
        "Dainaberry (Sleepberry)": "Dainaberry",
        "Devilweed (Wyssin)": "Devilweed",
        "Dungeon Fungus (Dungus)": "Dungeon Fungus",
        "Floure-du-luce (Blue Flag)": "Floure-du-luce",
        "Fumellar (Flower of Sleep)": "Fumellar",
        "Gotu Kola (Tiger's Calm)": "Gotu Kola",
        "Harrada Leaf (Barbarian's Boon)": "Harrada Leaf",
        "Hathil": "Hathlil",
        "Ipt (Roanwood)": "Ipt",
        "Jelly Moss (Ogre Snot)": "Jelly Moss",
        "Kiss of Discord (Lusiri Blossom)": "Kiss of Discord",
        "Kylathar (Changeberry)": "Kylathar",
        "Laishaberries (Fruit of Silence)": "Laishaberries",
        "Luurden (Bloodfruit)": "Luurden",
        "Maiden's Hair (Earthsilk)": "Maiden's Hair",
        "Obaddis Leaf\t": "Obaddis Leaf",
        "Psyllium (Plantain Weed)": "Psyllium",
        "Powdered Desert Milk (Wolves' Milk)": "Powdered Desert Milk",
        "Saracen's Confound": "Saracen's Compound",
        "Silverthorn (Aelebera)": "Silverthorn",
        "Sussur (Deeproot)": "Sussur",
        "Tahtoalethi (Wishfern)": "Tahtoalethi",
        "Tamariske (Blacktear)": "Tamariske",
        "Wildwood (Saelas)": "Wildwood",
        "Yarpick (Daggerthorn)": "Yarpick"
    }

    invalid_names = ['', '\n', '\n ']

    with open(plant_info_json) as json_file:
        data = json.load(json_file)

        # letters that have all normal entries
        for letter, entries in plants_by_letter.items():

            stupid_entry = None
            se_start = None
            se_len_extra_info = None
            letter_end = None

            if letter == 'A':
                stupid_entry = 'Alil'
                se_start = 151
                se_len_extra_info = 11
                letter_end = 296

            if letter == 'C':
                stupid_entry = 'Cow-Wheat'
                se_start = 548
                se_len_extra_info = 8
                letter_end = 580

            if letter == 'H':
                stupid_entry = 'Hangman Tree'
                se_start = 981
                se_len_extra_info = 12
                letter_end = 1068

            if letter == 'L':
                stupid_entry = 'Lizuara'
                se_start = 1237
                se_len_extra_info = 3
                letter_end = 1267

            if letter == 'M':
                stupid_entry = 'Marsh Maw'
                se_start = 1305
                se_len_extra_info = 13
                letter_end = 1398

            if letter == 'S':
                stupid_entry = 'Stygian Pumpkin'
                se_start = 1800
                se_len_extra_info = 3
                letter_end = 1836

            data = process_letter_entries(letter, entries, doc, invalid_names, inconsistencies, data,
                                          stupid_entry, se_start, se_len_extra_info, letter_end)

        with open(plant_info_v2_json, 'w') as outfile:
            json.dump(data, outfile, indent=4, sort_keys=True)


def generate_gmbinder_markdown():
    plant_info_json = OUTPUT_DIR / 'plant_info_v2.json'

    rarity_symbols = {
        'Very Common': 'VC',
        'Common': 'C',
        'Uncommon': 'U',
        'Rare': 'R',
        'Very Rare': 'VR',
        'Legendary': 'L',
        'Other': 'O'
    }

    plants_by_letter = collections.OrderedDict({
        'A': [], 'B': [], 'C': [], 'D': [], 'E': [], 'F': [], 'G': [], 'H': [], 'I': [], 'J': [],
        'K': [], 'L': [], 'M': [], 'N': [], 'O': [], 'P': [], 'Q': [], 'R': [], 'S': [], 'T': [],
        'U': [], 'V': [], 'W': [], 'X': [], 'Y': [], 'Z': []
    })
    plants_by_region = collections.OrderedDict({
        'Arctic': [],     # D&D 5E standard
        'City': [],
        'Coastal': [],    # D&D 5E standard - Coast
        'Desert': [],     # D&D 5E standard
        'Forest': [],     # D&D 5E standard
        'Jungle': [],
        'Plain': [],      # D&D 5E standard - Grassland
        'Mountain': [],   # D&D 5E standard
        'Ocean': [],
        'River': [],
        'Swamp': [],      # D&D 5E standard
        'Underdark': [],  # D&D 5E standard
        'Other': []
    })
    plants_by_rarity = collections.OrderedDict({
        'Very Common': [], 'Common': [], 'Uncommon': [], 'Rare': [], 'Very Rare': [], 'Legendary': [], 'Other': []
    })

    plants_for_table_entries = collections.OrderedDict()
    for region in plants_by_region.keys():
        plants_for_table_entries[region] = collections.OrderedDict()
        for rarity in plants_by_rarity.keys():
            plants_for_table_entries[region][rarity] = collections.OrderedDict({'die size': None,
                                                                                'plants': []})

    pages_before_plant_entries = 21
    header_height = 5  # equivalent # of lines for plant name, location, rarity
    footer_height = 1  # equivalent # of lines after entry
    lines_available_per_column = 60
    page_height = lines_available_per_column * 2
    desc_line_length = 60

    with open(plant_info_json) as json_file:
        data = json.load(json_file)

        pages = {}
        page_num = 1
        pages[page_num] = []
        current_page_height = 0

        for plant, entry in data.items():
            # plant = name
            # entry = {
            #   Regions: [x, y, z]
            #   Rarity: "xyz"
            #   Notes: ['' | 'Exceedingly Rare']
            #   Description: "foo"
            #   Extra info: ["this", "is", "extra"]
            #   ### added below ###
            #   name: "blah"
            #   page: 34
            #   first_letter: "B"
            #   homebrewery_height: 12
            #   rarity_symbol: "VC"
            #   table_die_entry: [None | 4 | 1-2]
            # }

            first_letter = plant[0].upper()
            description_height = math.ceil(len(entry['Description']) / desc_line_length)
            homebrewery_height = header_height + description_height + footer_height
            if plant == 'Alil':
                homebrewery_height += 40
            if plant == 'Cow-Wheat':
                homebrewery_height += 15
            if plant == 'Hangman Tree':
                homebrewery_height += 40
            if plant == 'Lizuara':
                homebrewery_height += 15
            if plant == 'Marsh Maw':
                homebrewery_height += 52
            if plant == 'Midnight Coneflower':
                homebrewery_height += 30
            if plant == 'Stygian Pumpkin':
                homebrewery_height += 15
            current_page_height += homebrewery_height

            if current_page_height > page_height:
                # print('page:', page_num, '   height:', current_page_height)
                current_page_height = homebrewery_height
                page_num += 1
                pages[page_num] = []

            entry['name'] = plant
            entry['page'] = page_num + pages_before_plant_entries
            entry['first_letter'] = first_letter
            entry['homebrewery_height'] = homebrewery_height
            entry['rarity_symbol'] = rarity_symbols[entry['Rarity']]
            entry['table_die_entry'] = None

            # Manually give descriptions line breaks to make content fit on page
            phrase = None
            if plant == 'Bodhi Tree':
                phrase = 'All the while'
            elif plant == 'Darkanda Bush':
                phrase = 'When ground up'
            elif plant == 'Elder Tree':
                phrase = 'Building a door'
            elif plant == 'Maraga Flowers':
                phrase = 'If the target creature'
            elif plant == 'Nettle':
                phrase = 'Aside from its'
            elif plant == 'Saddilia':
                phrase = 'The effect lasts'
            elif plant == 'Silver Tassel Toadstool':
                phrase = 'The poison is introduced'
            elif plant == 'Windwhip Tree':
                phrase = 'When the wind'

            if phrase:
                entry['Description'] = entry['Description'].replace(phrase, f'\n\n{phrase}')

            # manually deal with the 6 "special" plants
            if plant == 'Alil':
                entry['Extra info'] = '''| d10 | Psionic Ability |
|:----:|:-------------|
| 1  | Temporary intelligence bonus of +1. |
| 2  | Precision Mind: you become skilled at reading your foes. With this capability, you can call upon your gift to strike with increased accuracy. You add your Intelligence modifier (minimum 1) to an attack roll, after the roll, but before the announcement of the result. You can use this feature a number of times equal to your Intelligence modifier (minimum of 1). |
| 3  | Immunity to psychic damage. |
| 4  | You cannot be charmed or frightened for the duration. |
| 5  | You may take two actions during each of your turns for the duration. |
| 6  | You may add 1d6 psychic damage to any attack you make of 5 or less damage. |
| 7  | You project a field of improbability around yourself, creating a fleeting protective shell, gaining a +4 temporary bonus to AC for 3 rounds. Only usable once in a long-rest period. |
| 8  | You may, as a bonus action once per combat, instantly deliver a massive assault on the thought pathways of any one creature, dealing 1d10 points of psychic damage to it. |
| 9  | You heal another creature\u2019s wounds, transferring some of its damage to yourself. When you manifest this power as a full action in combat, you can heal as much as 2d10 points of damage to an ally by taking half of this damage to yourself. |
| 10 | You may gain proficiency in one extra skill until the long rest or proficiency with one tool or instrument for 1d6 days. |
'''

            if plant == 'Cow-Wheat':
                entry['Extra info'] = '''> ##### Drunkenness
> | Step | Effect |
> |:----:|:-------------|
> | 1  | -1 to Intelligence until sober. |
> | 2  | -1 to Wisdom until sober. |
> | 3  | -1 to Charisma until sober. |
> | 4  | -1 to Dexterity until sober. |
> | 5  | Disadvantage on all skill checks and attacks. |
> | 6  | Disadvantage on all saves. |
> | 7  | Blacks out for 1d4 hours. |
'''

            if plant == 'Hangman Tree':
                entry['Description'] = entry['Description'].replace('The main body', '\n\nThe main body')
                entry['Extra info'] = '''___
> ## Hangman Tree
>*Large tree, unaligned*
> ___
> - **Armor Class** 12
> - **Hit Points** 172(15d12 + 80)
> - **Speed** 1ft.
>___
>|STR|DEX|CON|INT|WIS|CHA|
>|:---:|:---:|:---:|:---:|:---:|:---:|
>|20 (+5)|6 (-2)|20 (+5)|3 (-4)|6 (-2)|4 (-3)|
>___
> - **Damage Resistances** bludgeoning, piercing
> - **Damage Vulnerabilities** fire
> - **Condition Immunities** blinded, deafened, frightened, exhaustion
> - **Senses** blindsight 90 ft., passive Perception 8
> - **Challenge** 8 (3,800 XP)
> ___
>
> ### Actions
> ***Multiattack.*** The hangman tree may make three *constrict* attacks per turn.
>
> ***Constrict.*** *Melee Weapon Attack:* +9 to hit, reach 10 ft., one target. *Hit* 10 (1d8 + 5) bludgeoning damage, and a Large or smaller target is grappled (escape DC 16).
            '''

            if plant in ['Lizuara', 'Stygian Pumpkin']:
                extra_info = ''
                for line in entry['Extra info']:
                    extra_info = extra_info + '\n\n' + line
                entry['Extra info'] = extra_info

            if plant == 'Marsh Maw':
                entry['Extra info'] = '''___
> ## Marsh Maw
>*Large plant, unaligned*
> ___
> - **Armor Class** 14
> - **Hit Points** 112 (14d10 + 28)
> - **Speed** 20 ft.. swim 20 ft.
>___
>|STR|DEX|CON|INT|WIS|CHA|
>|:---:|:---:|:---:|:---:|:---:|:---:|
>|16 (+3)|12 (+1)|14 (+2)|4 (-3)|10 (+0)|6 (-2)|
>___
> - **Damage Resistances** bludgeoning
> - **Damage Vulnerabilities** fire
> - **Condition Immunities** blinded, deafened, frightened, prone
> - **Senses** blindsight 60 ft. (blind beyond this radius), passive Perception 10
> - **Challenge** 6 (2,300 XP)
> ___
> ### Actions
> ***Multiattack.*** The marsh maw can make two *constrict* attacks and a *bite* attack.
>
> ***Constrict.*** *Melee Weapon Attack:* +5 to hit, reach 10 ft., one target. *Hit* 11 (2d6 + 3) bludgeoning damage, and a Large or smaller target is grappled (escape DC 14). Until this grapple ends, the target is restrained, and the marsh maw cannot constrict another target.
>
> ***Bite.*** *Melee Weapon Attack:* +5 to hit, reach 5 ft., one Medium or smaller target. Hit: 11 (2d6+3) piercing damage, and a target is blinded, restrained, and unable to breathe. The target must succeed on a DC 14 Con. Save at the start of each of the marsh maw’s turns or take 11 (2d8+3) bludgeoning damage. If the marsh maw moves, the engulfed target moves with it. The marsh maw is unable to use the bite attack until it releases the held creature.
'''

            if plant == 'Midnight Coneflower':
                entry['Extra info'] = '''If the powder of the Midnight Cone petals is combined with the Midnight Cone’s Nectar, the two can be mixed to create a sweet smelling, sweet tasting syrup that becomes inert when ingested. However, this process creates a powerful, delayed-response contact poison called Goodbye Kiss. Typical application sees the flower’s pure nectar applied to the assassin’s lips as a protective barrier followed by an application of the Goodbye Kiss syrup, and then the assassin giving the target a kiss. A creature subjected to this poison suffers no effect until the stroke of midnight. If the poison has not been neutralized before then, the creature must succeed on a DC 18 Constitution saving throw vs. poison, taking 13d6 poison damage on a failed save, or half as much damage on a successful one.

In addition, the Midnight Cone Flower’s petals can be made into Midnight Tears, an equally potent ingested poison. Many nobles have died peacefully in their sleep while their loved ones lay nearby, falling to the effects of a dose of Midnight Tears administered by a stealthy assassin many hours before. The only side-effect of the poison, that can very easily be missed in investigation, is the presence of the white, pearlescent tears that the poison is named for, seeping from the recipient’s eyes as if they were crying in their sleep. Again, a creature subjected to this poison suffers no effect until the stroke of midnight. If the poison has not been neutralized before then, the creature must succeed on a DC 18 Constitution saving throw vs. poison, taking 13d6 poison damage on a failed save, or half as much damage on a successful one.'''

            pages[page_num].append(entry)
            plants_by_letter[first_letter].append(entry)
            plants_by_rarity[entry['Rarity']].append(entry)
            for region in entry['Regions']:
                plants_by_region[region].append(entry)
                plants_for_table_entries[region][entry['Rarity']]['plants'].append(entry)

        plants_by_letter = collections.OrderedDict(sorted(plants_by_letter.items(), key=lambda t: len(t)))
        plants_by_rarity = collections.OrderedDict(sorted(plants_by_rarity.items(), key=lambda t: len(t)))
        plants_by_region = collections.OrderedDict(sorted(plants_by_region.items(), key=lambda t: len(t)))

        # second pass clean-up for table entries
        for region, rarity_dict in plants_for_table_entries.items():
            for rarity, plant_info_dict in rarity_dict.items():
                num_plants = len(plants_for_table_entries[region][rarity]['plants'])
                die_size = find_die_size(num_plants)
                plants_for_table_entries[region][rarity]['die size'] = die_size

                die_entries = get_die_entries(num_plants, die_size)

                for x in range(num_plants):
                    plants_for_table_entries[region][rarity]['plants'][x] = plants_for_table_entries[region][rarity]['plants'][x].copy()
                    plants_for_table_entries[region][rarity]['plants'][x]['table_die_entry'] = die_entries[x]

        # just printing some debug stuff here
        # for region, region_entries in plants_by_region.items():
        #     print(f'{region} ({len(region_entries)})')
        #     for rarity, entries in plants_by_rarity.items():
        #         plant_table_entries = [x for x in plants_by_rarity[rarity] if region in x['Regions']]
        #         die_size = plants_for_table_entries[region][rarity]["die size"]
        #         print(f'    {rarity}: {len(plant_table_entries)} (d{die_size})')

        index_pages = {
            'Arctic': '107',
            'City': '107',
            'Coastal': '107',
            'Desert': '107',
            'Forest': '107',
            'Jungle': '108',
            'Plain': '109',
            'Mountain': '110',
            'Ocean': '111',
            'River': '111',
            'Swamp': '111',
            'Underdark': '111',
            'Other': '112',
            'Very Common': '113',
            'Common': '113',
            'Uncommon': '114',
            'Rare': '115',
            'Very Rare': '115',
            'Legendary': '116'
        }

        context = {
            'title': 'Broderick’s Compendium: Fantasy Plants Across the Realms',
            'plants_by_letter': plants_by_letter,
            'plants_by_region': plants_by_region,
            'plants_by_rarity': plants_by_rarity,
            'plants_for_table_entries': plants_for_table_entries,
            'pages': pages,
            'index_pages': index_pages
        }

        template = TEMPLATE_DIR / 'gmbinder.md.j2'
        generated_file = OUTPUT_DIR / "generated-gmbinder.txt"

        pages_dir = OUTPUT_DIR / 'pages'
        pages_dir.mkdir(parents=True, exist_ok=True)

        result = render(template, context)
        pages = result.split('\\pagebreakNum')
        pages_per_file = 10
        page_counter = 1

        with codecs.open(TEMPLATE_DIR / 'gmbinder-css.md.j2', "r", encoding="utf-8") as css_file:
            css_data = css_file.read()

            for page_list in chunker(pages, pages_per_file):
                end_page = page_counter + pages_per_file
                page_file = OUTPUT_DIR / 'pages' / f'pages-{page_counter}-{min([end_page, len(pages)])}-raw.txt'

                with codecs.open(page_file, "w", encoding="utf-8") as _file:
                    if page_counter != 1:
                        _file.write(css_data)
                    for page_string in page_list:
                        _file.write(page_string)
                        if page_string != page_list[-1]:
                            _file.writelines('\n\\pagebreakNum\n')
                page_counter += pages_per_file

        with codecs.open(generated_file, "w", encoding="utf-8") as _file:
            _file.write(result)


def chunker(seq, size):
    return (seq[pos:pos + size] for pos in range(0, len(seq), size))


def find_die_size(num_plants):
    die_sizes = [0, 4, 6, 8, 10, 12, 20, 100]
    for die_size in die_sizes:
        if num_plants <= die_size:
            return die_size


def get_die_entries(num_plants, die_size):
    # die_entries = Queue()

    die_entries = []
    table_increments = get_table_increments(num_plants, die_size)
    # print(table_increments)

    if num_plants == die_size:
        die_entries = [x + 1 for x in range(num_plants)]
    elif num_plants < die_size:
        previous_high = 0
        for x in range(num_plants):
            inc = table_increments[x]
            start = previous_high + 1
            end = start + inc
            if start == end:
                die_entries.append(start)
            else:
                die_entries.append(f'{start}-{end}')
            previous_high = end

    else:
        # print(f'num_plants ({num_plants}) needs to be less than die size ({die_size})!')
        return None

    return die_entries


def get_table_increments(num_plants, die_size):
    foo = [0 for _ in range(num_plants)]
    if num_plants == die_size:
        return foo
    target = die_size - num_plants
    counter = 0
    for x in foo:
        foo[counter] = x + math.ceil(target / num_plants)
        counter += 1
        target -= 1
        if target == 0:
            break
    return foo


def render(template, context):
    path, filename = os.path.split(template)
    env = jinja2.Environment(loader=jinja2.FileSystemLoader(path or './'))
    env.trim_blocks = True
    env.lstrip_blocks = True
    return env.get_template(filename).render(context)


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    # deduplicate_csv()
    # convert_csv_to_json()
    # parse_description()
    generate_gmbinder_markdown()


if __name__ == "__main__":
    main()
